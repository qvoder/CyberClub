# -*- coding: utf-8 -*-
"""Добавляет в пояснительную записку раздел про страницу магазина."""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_block_after(paragraph, lines, style=None):
    last = paragraph
    for line in lines:
        if line == '':
            last = insert_paragraph_after(last, '', style)
        else:
            last = insert_paragraph_after(last, line, style)
    return last


def main():
    path = r'E:\WorkSpace\CyberClub\report\Пояснительная записка.docx'
    doc = Document(path)

    for p in doc.paragraphs:
        t = p.text
        if t.startswith('Сайт «Кибертека» построен по иерархическому принципу'):
            if 'shop.html' not in t:
                p.add_run(
                    ' Дополнительно реализована отдельная страница онлайн-магазина (shop.html) — '
                    'каталог снеков, напитков и горячих блюд с корзиной и выбором филиала для получения заказа. '
                    'Страница доступна из бокового меню на всех основных разделах сайта.'
                )
        elif t.startswith('На сайте реализовано два типа навигации'):
            p.text = (
                'На сайте реализовано два типа навигации: основное вертикальное боковое меню (сайдбар) '
                'и навигация через плитки (карточки зон). Боковое меню присутствует на внутренних страницах '
                '(profsouz.html, bauman.html, profprice.html, profbauman.html, shop.html и др.). '
                'Стили меню вынесены в общий файл css/main-nav.css (компактные размеры под семь пунктов: '
                'ширина полосы 88 px, уменьшенные межстрочные интервалы). '
                'Меню зафиксировано слева (position: fixed) и содержит ссылки: «Главная», «Зоны», «Цены», '
                '«Фото», «Акции», «Магазин», «Контакты». Ссылка «Магазин» ведёт на shop.html. '
                'Ссылка «Цены» ведёт на страницу тарифов. При прокрутке скрипт navigation.js подсвечивает '
                'активный пункт (на страницах с якорными секциями). На мобильных устройствах меню скрывается. '
                'Пример HTML-кода меню:'
            )
        elif '#akcii' in t and 'main-nav__link' in t and 'Магазин' not in t:
            insert_paragraph_after(
                p,
                '            <a href="shop.html" class="main-nav__link">Магазин</a>',
            )

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith('В данном проекте были реализованы следующие приёмы'):
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            if 'CeoOptimization.js' in p.text and p.text.startswith('Скрипт SEO'):
                anchor = p
                for _ in range(8):
                    n = anchor._p.getnext()
                    if n is None:
                        break
                    anchor = Paragraph(n, anchor._parent)
                break

    if anchor is None:
        for p in reversed(doc.paragraphs):
            if p.text.strip() == 'ЗАКЛЮЧЕНИЕ':
                anchor = doc.paragraphs[doc.paragraphs.index(p) - 1]
                break

    block = [
        '',
        'Страница онлайн-магазина (shop.html)',
        'В рамках расширения функциональности сайта добавлена отдельная страница магазина shop.html. '
        'Она оформлена в единой дизайн-системе «Кибертека» (шрифт Inter, тёмный фон, акцентные цвета розовый/синий/красный) '
        'и подключает файлы: css/main-nav.css, css/shop.css, библиотеку Swiper 11 (CDN), скрипт js/shop.js.',
        'Структура страницы: прозрачное боковое меню; блок hero с логотипом и заголовком «Магазин Кибертеки»; '
        'кнопка «Корзина»; каталог из трёх категорий (снеки, напитки, горячие блюда); выезжающая панель корзины; '
        'модальное окно карточки товара; модальное окно симуляции оплаты.',
        'Каталог товаров хранится во внешнем JSON-файле data/shop-products.json (45 позиций: по 15 в каждой категории). '
        'Для каждого товара заданы поля: id, category, name, price, description, details, weight, image. '
        'Путь в поле image указывается относительно корня сайта, например img/shop/snacks-01.jpg; пустая строка — '
        'заглушка с логотипом клуба. Инструкция по добавлению фотографий приведена в файле data/КАК-ДОБАВИТЬ-ФОТО.txt.',
        'На экранах шире 768 px товары выводятся CSS-сеткой (grid). На мобильных устройствах для каждой категории '
        'используется горизонтальный слайдер Swiper — все позиции прокручиваются влево-вправо, без длинной колонки.',
        'Подключение скриптов и стилей на странице магазина:',
        '<link rel="stylesheet" href="css/main-nav.css">',
        '<link rel="stylesheet" href="css/shop.css">',
        '<script src="https://cdn.jsdelivr.net/npm/swiper@11/swiper-bundle.min.js"></script>',
        '<script src="js/shop.js"></script>',
        '',
        'Скрипт магазина (shop.js)',
        'Файл js/shop.js реализует загрузку каталога, корзину, выбор филиала и работу интерфейса без перезагрузки страницы.',
        'Основные функции:',
        '— fetch(\'data/shop-products.json\') — загрузка товаров при открытии страницы (требуется локальный HTTP-сервер);',
        '— отображение карточек по категориям; на карточке кнопка «+», после добавления — «−», счётчик и «+»;',
        '— клик по карточке открывает модальное окно с фото, описанием и ценой из JSON;',
        '— корзина: изменение количества, выбор клуба (Бауманская / Профсоюзная), сохранение в localStorage;',
        '— симуляция оплаты (проверка полей формы, очистка корзины после «успешной» оплаты).',
        'Фрагмент логики счётчика на карточке:',
        'function setCartQty(productId, qty) {',
        '    if (qty <= 0) delete cartState.items[productId];',
        '    else cartState.items[productId] = qty;',
        '    saveCart();',
        '    updateAllCardQtyUI();',
        '    renderCart();',
        '}',
        '',
        'Стили магазина (shop.css)',
        'Файл css/shop.css задаёт оформление hero-блока (градиент и фон как на главной), сетку и карточки товаров, '
        'панель корзины, модальные окна, кнопки в стиле сайта. Для мобильной вёрстки скрывается сетка и показывается '
        'блок .shop-catalog__slider с навигационными стрелками.',
        '',
        'Файл css/main-nav.css',
        'Общие стили бокового меню для страниц филиалов, тарифов и магазина: переменные --nav-rail-w и --nav-content-pad, '
        'единые размеры шрифта и отступы контента под увеличенное число пунктов меню.',
        '',
        'Рисунок – страница онлайн-магазина (shop.html)',
        '',
        'Листинг фрагмента shop.html (подключение и навигация):',
        '<!DOCTYPE html>',
        '<html lang="ru">',
        '<head>',
        '    <link rel="stylesheet" href="css/main-nav.css">',
        '    <link rel="stylesheet" href="css/shop.css">',
        '</head>',
        '<body class="shop-page">',
        '<nav class="main-nav">',
        '    <a href="index.html" class="main-nav__link">Главная</a>',
        '    …',
        '    <a href="shop.html" class="main-nav__link shop-nav-active">Магазин</a>',
        '    <a href="bauman.html#kontakty" class="main-nav__link">Контакты</a>',
        '</nav>',
        '<main class="shop-main" id="catalog"></main>',
        '<script src="js/shop.js"></script>',
        '</body>',
        '</html>',
    ]

    if anchor:
        insert_block_after(anchor, block)

    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == 'ЗАКЛЮЧЕНИЕ' and i > 0:
            prev = paras[i - 1]
            if 'магазин' not in prev.text.lower() and 'shop.html' not in prev.text:
                insert_paragraph_after(
                    prev,
                    'Дополнительно на сайте реализована страница онлайн-магазина (shop.html) с каталогом из JSON, '
                    'корзиной и адаптивными горизонтальными слайдерами категорий на мобильных устройствах, '
                    'что расширяет сервис клуба для гостей.',
                )
            break

    for p in doc.paragraphs:
        if p.text.strip() == 'booking-popup.css (модальное окно)':
            insert_block_after(
                p,
                [
                    'shop.css (страница магазина)',
                    'css',
                    '.shop-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(200px, 1fr)); }',
                    '.shop-card--in-cart { border-color: rgba(230, 55, 183, 0.35); }',
                    '@media (max-width: 768px) { .shop-catalog__grid { display: none; } .shop-catalog__slider { display: block; } }',
                    'main-nav.css (общее боковое меню)',
                    'css',
                    ':root { --nav-rail-w: 88px; --nav-content-pad: calc(var(--nav-rail-w) + 12px); }',
                    '.main-nav__list { gap: 28px; }',
                    '.main-nav__link { font-size: 22px; }',
                ],
            )
            break

    doc.save(path)
    print('OK: updated', path)


if __name__ == '__main__':
    main()
